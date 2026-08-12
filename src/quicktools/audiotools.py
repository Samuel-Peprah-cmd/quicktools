"""Audio transcription utilities, powered by faster-whisper (an optimized implementation
of OpenAI's Whisper speech recognition model) and Pyannote for Speaker Diarization.

Requires the optional 'faster-whisper' package: pip install faster-whisper
Diarization requires 'pyannote.audio': pip install pyannote.audio
The first call for a given model_size downloads that model (requires internet, one-time).
No separate ffmpeg installation is required — audio decoding is bundled in.
"""
import os
import tempfile
import warnings
warnings.filterwarnings("ignore", message=".*torchcodec.*")
warnings.filterwarnings("ignore", category=UserWarning, module="pyannote.*")

def _load_model(model_size: str = "base", device: str = "cpu"):
    """
    Initializes the faster-whisper model and automatically assigns compute type
    based on the available device (GPU vs CPU).
    """
    try:
        from faster_whisper import WhisperModel
    except ImportError:
        raise ImportError(
            "Audio transcription requires faster-whisper. Install it with: pip install faster-whisper"
        )
    
    # Map 'auto' or 'cuda' for GPU hardware acceleration
    compute_type = "float16" if device == "cuda" else "int8"
    
    return WhisperModel(model_size, device=device, compute_type=compute_type)


def transcribe_audio(path_or_url: str, model_size: str = "base", language: str | None = None, task: str = "transcribe") -> str:
    """Transcribe local audio/video or web URLs to plain text. Supports task='translate' to convert to English."""
    import os
    import time
    
    temp_file = None
    if path_or_url.startswith("http://") or path_or_url.startswith("https://"):
        audio_path = download_url_audio(path_or_url)
        temp_file = audio_path
    else:
        audio_path = path_or_url

    try:
        # Prevent crash if frontend sends 'auto' string instead of None
        if language == "auto" or language == "":
            language = None

        # Determine if we should route to Gemini
        lang_lower = language.lower() if language else ""
        use_gemini = lang_lower in ["twi", "akan", "en", "english"]

        # === AI ENGINE (Powered by Gemini 3.6 Flash) ===
        if use_gemini:
            try:
                from google import genai
            except ImportError:
                raise RuntimeError("The 'google-genai' package is missing. Please install it.")
                
            api_key = os.environ.get("GEMINI_API_KEY")
            if not api_key:
                raise RuntimeError("GEMINI_API_KEY is missing from your environment variables!")

            print(f"  [AtomDev] Intercepted {lang_lower} request. Routing audio to Gemini 3.6 Engine...")
            client = genai.Client(api_key=api_key)
            
            # Upload the audio file directly to Gemini's memory
            print("  [AtomDev] Uploading media to Gemini...")
            audio_file = client.files.upload(file=audio_path)
            
            # Dynamically set the prompt based on language and task
            if lang_lower in ["twi", "akan"]:
                if task == "translate":
                    print("  [AtomDev] Translating mixed Twi/English directly to English...")
                    prompt = (
                        "Listen to this audio, which may contain a mix of Twi/Akan and English. "
                        "Translate any Twi spoken into highly accurate, natural-sounding English, "
                        "and accurately transcribe the English parts. Provide a clean, well-formatted, "
                        "and punctuated transcript. Remove verbal stutters and false starts. "
                        "Organize the dialogue clearly with plain-text speaker labels in ALL CAPS "
                        "(e.g., INTERVIEWER:, RESPONDENT:). Strictly DO NOT use asterisks or any markdown formatting. "
                        "Return ONLY the final unified English text."
                    )
                else:
                    print("  [AtomDev] Transcribing Twi...")
                    prompt = (
                        "Listen to this Twi/Akan audio. Provide a highly accurate transcription in Twi. "
                        "Provide a clean, well-formatted, and punctuated transcript. "
                        "Organize the dialogue clearly with plain-text speaker labels in ALL CAPS "
                        "(e.g., INTERVIEWER:, RESPONDENT:). Strictly DO NOT use asterisks or any markdown formatting. "
                        "Return ONLY the finalized Twi text."
                    )
            else:
                # English Cleanup Prompt
                print("  [AtomDev] Transcribing and cleaning English audio...")
                prompt = (
                    "Listen to this English interview audio. Provide a clean, well-formatted, and punctuated transcript. "
                    "Remove verbal stutters, false starts, filler words, and background chatter. "
                    "Organize the dialogue clearly with plain-text speaker labels in ALL CAPS "
                    "(e.g., INTERVIEWER:, RESPONDENT:). Strictly DO NOT use asterisks or any markdown formatting. "
                    "Return ONLY the cleaned transcript text."
                )
                
            try:
                # Generate text with an automatic retry loop for 503 errors
                max_retries = 3
                final_text = ""
                
                for attempt in range(max_retries):
                    try:
                        response = client.models.generate_content(
                            model='gemini-3.6-flash',
                            contents=[prompt, audio_file]
                        )
                        final_text = response.text.strip()
                        break  # Success! Break out of the retry loop.
                        
                    except Exception as e:
                        error_msg = str(e)
                        if "503" in error_msg or "UNAVAILABLE" in error_msg:
                            if attempt < max_retries - 1:
                                wait_time = 3 * (attempt + 1)
                                print(f"  [AtomDev] Google servers busy (503). Retrying in {wait_time} seconds...")
                                time.sleep(wait_time)
                            else:
                                raise RuntimeError("Google Gemini servers are currently overloaded. Please try again in a few minutes.")
                        else:
                            # If it's a different error (like 400 or 403), raise it immediately
                            raise e
            finally:
                # Clean up the file from cloud memory
                try:
                    client.files.delete(name=audio_file.name)
                except Exception:
                    pass
                
            return final_text
        # ======================================================

        # === STANDARD ENGINE (For all other languages, e.g., French, Spanish) ===
        import torch
        from faster_whisper import WhisperModel
        
        device_str = "cuda" if torch.cuda.is_available() else "cpu"
        compute_type = "float16" if device_str == "cuda" else "int8"
        
        print(f"  [AtomDev] Processing with standard {model_size} engine...")
        model = WhisperModel(model_size, device=device_str, compute_type=compute_type)
        segments, _ = model.transcribe(audio_path, language=language, task=task)
        return " ".join(segment.text.strip() for segment in segments)

    finally:
        if temp_file and os.path.exists(temp_file):
            import shutil
            shutil.rmtree(os.path.dirname(temp_file), ignore_errors=True)

def transcribe_audio_with_timestamps(path: str, model_size: str = "base") -> list[dict]:
    """Transcribe audio into a list of segments, each with 'start', 'end' (seconds), and 'text'."""
    model = _load_model(model_size)
    segments, _ = model.transcribe(path)
    return [{"start": s.start, "end": s.end, "text": s.text.strip()} for s in segments]


def transcribe_audio_word_level(path: str, model_size: str = "base") -> list[dict]:
    """Transcribe audio into a word-by-word list, each with 'word', 'start', and 'end' (seconds).
    This gives true word-for-word timing, useful for captions or karaoke-style highlighting."""
    model = _load_model(model_size)
    segments, _ = model.transcribe(path, word_timestamps=True)
    words = []
    for segment in segments:
        for word in segment.words:
            words.append({"word": word.word.strip(), "start": word.start, "end": word.end})
    return words


def detect_audio_language(path: str, model_size: str = "base") -> str:
    """Detect the spoken language of an audio file, returning its ISO language code (e.g. 'en')."""
    model = _load_model(model_size)
    _, info = model.transcribe(path)
    return info.language


def save_transcript_as_srt(segments: list[dict], output_path: str) -> None:
    """Save a list of transcript segments (from transcribe_audio_with_timestamps) as an .srt
    subtitle file, ready to use with any video player or editor."""
    def format_timestamp(seconds: float) -> str:
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds - int(seconds)) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

    lines = []
    for i, seg in enumerate(segments, start=1):
        lines.append(str(i))
        lines.append(f"{format_timestamp(seg['start'])} --> {format_timestamp(seg['end'])}")
        lines.append(seg["text"])
        lines.append("")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


# --- NEW DIARIZATION AND STYLING FEATURES ---

def _load_waveform_via_av(path: str, target_sr: int = 16000):
    """Decode any audio/video file into a mono waveform tensor using PyAV — bypassing
    pyannote's own (fragile, Windows-unfriendly) torchcodec-based audio loader entirely."""
    import av
    import torch
    import numpy as np

    container = av.open(path)
    resampler = av.AudioResampler(format="s16", layout="mono", rate=target_sr)
    audio_stream = container.streams.audio[0]

    chunks = []
    for frame in container.decode(audio_stream):
        for resampled in resampler.resample(frame):
            chunks.append(resampled.to_ndarray())
    container.close()

    if not chunks:
        raise RuntimeError(f"No audio could be decoded from '{path}'")

    audio = np.concatenate(chunks, axis=1).astype(np.float32) / 32768.0  # int16 -> [-1, 1]
    waveform = torch.from_numpy(audio)  # shape: (1, num_samples) since layout="mono"
    return waveform, target_sr


def download_url_audio(url: str) -> str:
    """Download audio stream from YouTube, TikTok, Instagram, X, etc. via yt-dlp."""
    try:
        import yt_dlp
    except ImportError:
        raise ImportError(
            "Downloading audio from URLs requires 'yt-dlp'. Install it with: pip install yt-dlp"
        )

    temp_dir = tempfile.gettempdir()
    ydl_opts = {
        'format': 'bestaudio/best',
        'outtmpl': os.path.join(temp_dir, 'quicktools_%(id)s.%(ext)s'),
        'quiet': True,
        'no_warnings': True,
    }
    print(f"📥 Fetching audio stream from URL: {url}...")
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=True)
        filename = ydl.prepare_filename(info)
        return filename


def _get_diarization_turns(output):
    """Extract diarization turns handling pyannote.audio 3.x and 4.x."""
    if hasattr(output, "speaker_diarization"):
        diarization = output.speaker_diarization
    else:
        diarization = output

    turns = []
    for turn, _, speaker in diarization.itertracks(yield_label=True):
        turns.append((turn, speaker))
    return turns

def chunk_audio_file(input_path: str, chunk_duration_ms: int = 600000) -> list[str]:
    """
    Splits a massive audio file into smaller segments (default 10 minutes / 600,000ms) 
    to prevent memory exhaustion on server nodes.
    
    Returns a list of temporary file paths for the chunks.
    """
    try:
        import av
    except ImportError:
        raise ImportError("Chunking requires PyAV. Install it with: pip install av")
    
    container = av.open(input_path)
    audio_stream = container.streams.audio[0]
    
    temp_dir = tempfile.mkdtemp()
    chunk_paths = []
    
    chunk_idx = 0
    output_path = os.path.join(temp_dir, f"chunk_{chunk_idx}.wav")
    
    output_container = av.open(output_path, mode="w", format="wav")
    output_stream = output_container.add_stream("pcm_s16le", rate=16000, layout="mono")
    
    start_time = 0.0
    for frame in container.decode(audio_stream):
        if frame.time - start_time > (chunk_duration_ms / 1000.0):
            output_container.close()
            chunk_paths.append(output_path)
            
            chunk_idx += 1
            output_path = os.path.join(temp_dir, f"chunk_{chunk_idx}.wav")
            output_container = av.open(output_path, mode="w", format="wav")
            output_stream = output_container.add_stream("pcm_s16le", rate=16000, layout="mono")
            start_time = frame.time
            
        for packet in output_stream.encode(frame):
            output_container.mux(packet)
            
    output_container.close()
    container.close()
    chunk_paths.append(output_path)
    
    return chunk_paths


def transcribe_with_speakers(path_or_url: str, hf_token: str, model_size: str = "base",
                             device: str = "auto", language: str | None = None) -> list[dict]:
    """Transcribe local audio/video or web URLs (YouTube, TikTok, IG, X) with word-level speaker diarization."""
    temp_file = None
    
    # Auto-detect if input is a Web URL
    if path_or_url.startswith("http://") or path_or_url.startswith("https://"):
        temp_file = download_url_audio(path_or_url)
        audio_path = temp_file
    else:
        audio_path = path_or_url

    try:
        if os.name == 'nt':
            print("💡 Tip: For faster native audio processing, you can install FFmpeg.")
            print("   Run this command in PowerShell: winget install ffmpeg\n")
            
        try:
            from pyannote.audio import Pipeline
        except ImportError:
            raise ImportError(
                "Speaker diarization requires pyannote.audio. Install with: pip install pyannote.audio"
            )

        import torch
        if device == "auto":
            device = "cuda" if torch.cuda.is_available() else "cpu"

        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)

            print("[1/4] Decoding audio stream...")
            waveform, sample_rate = _load_waveform_via_av(audio_path)

            print("[2/4] Loading diarization model...")
            pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization-3.1", token=hf_token)
            
            if device == "cuda":
                pipeline.to(torch.device("cuda"))
                print("⚡ NVIDIA GPU enabled.")
            else:
                print("🐢 Running diarization on CPU.")

            print("[3/4] Running speaker diarization...")
            try:
                from pyannote.audio.pipelines.utils.hook import ProgressHook
                with ProgressHook() as hook:
                    diarization = pipeline({"waveform": waveform, "sample_rate": sample_rate}, hook=hook)
            except ImportError:
                diarization = pipeline({"waveform": waveform, "sample_rate": sample_rate})

        print("[4/4] Transcribing speech with word-level timestamps...")
        
        # Prevent crash if frontend sends 'auto' string instead of None
        if language == "auto" or language == "":
            language = None
            
        model = _load_model(model_size)
        segments, _ = model.transcribe(audio_path, word_timestamps=True, language=language)

        # Extract individual words with exact timestamps
        all_words = []
        for segment in segments:
            if hasattr(segment, 'words') and segment.words:
                for word in segment.words:
                    all_words.append(word)
            else:
                all_words.append(segment)

        turns = _get_diarization_turns(diarization)

        # Map each word to Pyannote speaker turns
        word_speaker_map = []
        for w in all_words:
            w_mid = (w.start + w.end) / 2
            speaker = "Unknown"
            for turn, spk in turns:
                if turn.start <= w_mid <= turn.end:
                    speaker = spk
                    break
            text_content = getattr(w, 'word', getattr(w, 'text', ''))
            word_speaker_map.append({
                "start": w.start,
                "end": w.end,
                "speaker": speaker.replace("SPEAKER_", "Speaker "),
                "word": text_content
            })

        # Re-group consecutive words belonging to the same speaker
        transcript_data = []
        if word_speaker_map:
            curr_speaker = word_speaker_map[0]["speaker"]
            curr_start = word_speaker_map[0]["start"]
            curr_end = word_speaker_map[0]["end"]
            curr_words = [word_speaker_map[0]["word"]]

            for w in word_speaker_map[1:]:
                if w["speaker"] == curr_speaker:
                    curr_words.append(w["word"])
                    curr_end = w["end"]
                else:
                    transcript_data.append({
                        "start": round(curr_start, 2),
                        "end": round(curr_end, 2),
                        "speaker": curr_speaker,
                        "text": "".join(curr_words).strip()
                    })
                    curr_speaker = w["speaker"]
                    curr_start = w["start"]
                    curr_end = w["end"]
                    curr_words = [w["word"]]

            transcript_data.append({
                "start": round(curr_start, 2),
                "end": round(curr_end, 2),
                "speaker": curr_speaker,
                "text": "".join(curr_words).strip()
            })

        print("Done!")
        return transcript_data

    finally:
        if temp_file and os.path.exists(temp_file):
            try:
                os.remove(temp_file)
            except Exception:
                pass

def save_transcript_to_docx(transcript_data: list[dict], output_path: str) -> None:
    """
    Takes speaker-mapped transcript data and generates a beautifully styled Word document.
    Requires the 'python-docx' package.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise ImportError("Saving to Word requires python-docx. Run: pip install python-docx")
        
    doc = Document()
    title = doc.add_heading("Audio Meeting Transcript", level=1)
    title.alignment = 1 # Center align

    # Track the last speaker so we don't print the name over and over for continuous talking
    last_speaker = None

    for entry in transcript_data:
        p = doc.add_paragraph()
        
        # Convert raw seconds into [MM:SS] format
        start_m, start_s = divmod(int(entry['start']), 60)
        time_str = f"[{start_m:02d}:{start_s:02d}]"
        
        # Only print the Speaker ID if the speaker changed
        if entry['speaker'] != last_speaker:
            speaker_run = p.add_run(f"{time_str} {entry['speaker']}:\n")
            speaker_run.bold = True
            # Give the speaker tag a nice professional blue color
            speaker_run.font.color.rgb = RGBColor(0, 102, 204) 
            last_speaker = entry['speaker']
        
        # Add the actual spoken text
        p.add_run(entry['text'])
        p.paragraph_format.space_after = Pt(8)
        
    doc.save(output_path)