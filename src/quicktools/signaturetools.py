"""Digital signature utilities: visual signature stamping onto PDFs/Word docs, and
true cryptographic PDF signing (PAdES-style) so a document carries both a professional
hand-drawn appearance and tamper-evident cryptographic proof.

Visual stamping (stamp_signature_on_pdf, stamp_signature_on_docx) has no extra
dependencies beyond what quicktools already uses.

Certificate generation (generate_signing_certificate) requires the 'cryptography'
package (already a transitive dependency of several quicktools modules).

True cryptographic signing (sign_pdf_cryptographically) requires the optional
'pyhanko' package: pip install pyhanko
"""
import io


def stamp_signature_on_pdf(pdf_path: str, signature_image_path: str, output_path: str,
                            page_number: int = -1, x: float = 72, y: float = 72,
                            width: float = 200, height: float = 80) -> None:
    """Stamp a (typically transparent-background) signature image onto a PDF page.
    page_number is 1-indexed; use -1 (default) for the last page. x/y/width/height
    are in points, measured from the bottom-left of the page."""
    from pypdf import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas

    reader = PdfReader(pdf_path)
    writer = PdfWriter()

    target_index = len(reader.pages) - 1 if page_number == -1 else page_number - 1
    if not (0 <= target_index < len(reader.pages)):
        raise ValueError(f"page_number {page_number} is out of range for a {len(reader.pages)}-page PDF")

    for i, page in enumerate(reader.pages):
        if i == target_index:
            page_width = float(page.mediabox.width)
            page_height = float(page.mediabox.height)

            packet = io.BytesIO()
            c = canvas.Canvas(packet, pagesize=(page_width, page_height))
            c.drawImage(signature_image_path, x, y, width=width, height=height,
                        mask="auto", preserveAspectRatio=True)
            c.save()
            packet.seek(0)

            overlay_reader = PdfReader(packet)
            page.merge_page(overlay_reader.pages[0])
        writer.add_page(page)

    with open(output_path, "wb") as f:
        writer.write(f)


def stamp_signature_on_docx(docx_path: str, signature_image_path: str, output_path: str,
                            replace_placeholder: str | None = None, width_inches: float = 2.0) -> None:
    import docx
    from docx.shared import Inches

    doc = docx.Document(docx_path)
    inserted = False

    def replace_in_paragraph(p):
        if replace_placeholder and replace_placeholder in p.text:
            # Split the paragraph text around the placeholder to keep the underscores/lines
            parts = p.text.split(replace_placeholder)
            
            # Clear all existing hidden runs in this paragraph
            for r in p.runs:
                r.text = ""
                
            # Rebuild the paragraph with the image injected perfectly in the middle
            if parts[0]:
                p.add_run(parts[0])
                
            run = p.add_run()
            run.add_picture(signature_image_path, width=Inches(width_inches))
            
            if len(parts) > 1 and parts[1]:
                p.add_run(parts[1])
                
            return True
        return False

    # 1. Search all standard paragraphs
    if replace_placeholder:
        for p in doc.paragraphs:
            if replace_in_paragraph(p):
                inserted = True
                break  # Stop after placing the first signature

        # 2. Search inside tables (signatures are very often placed in document tables)
        if not inserted:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if replace_in_paragraph(p):
                                inserted = True
                                break
                        if inserted: break
                    if inserted: break

    # 3. Fallback: If no placeholder was provided or found, append to the very end
    if not inserted:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(signature_image_path, width=Inches(width_inches))

    doc.save(output_path)

def generate_signing_certificate(common_name: str, output_cert_path: str, output_key_path: str,
                                  valid_days: int = 365, key_password: str | None = None) -> None:
    """Generate a self-signed X.509 certificate and private key pair, suitable for
    cryptographically signing PDFs (see sign_pdf_cryptographically). For a signature
    that browsers/Adobe Acrobat show as fully 'trusted' rather than 'self-signed',
    you'd instead obtain a certificate from a recognized Certificate Authority —
    this self-signed version still provides genuine tamper-evidence, just without
    third-party identity verification."""
    from cryptography import x509
    from cryptography.x509.oid import NameOID
    from cryptography.hazmat.primitives import hashes, serialization
    from cryptography.hazmat.primitives.asymmetric import rsa
    from datetime import datetime, timedelta, timezone

    private_key = rsa.generate_private_key(public_exponent=65537, key_size=2048)

    subject = issuer = x509.Name([
        x509.NameAttribute(NameOID.COMMON_NAME, common_name),
    ])

    cert = (
        x509.CertificateBuilder()
        .subject_name(subject)
        .issuer_name(issuer)
        .public_key(private_key.public_key())
        .serial_number(x509.random_serial_number())
        .not_valid_before(datetime.now(timezone.utc))
        .not_valid_after(datetime.now(timezone.utc) + timedelta(days=valid_days))
        .sign(private_key, hashes.SHA256())
    )

    encryption = (
        serialization.BestAvailableEncryption(key_password.encode())
        if key_password else serialization.NoEncryption()
    )

    with open(output_key_path, "wb") as f:
        f.write(private_key.private_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PrivateFormat.PKCS8,
            encryption_algorithm=encryption,
        ))

    with open(output_cert_path, "wb") as f:
        f.write(cert.public_bytes(serialization.Encoding.PEM))


def sign_pdf_cryptographically(pdf_path: str, output_path: str, cert_path: str, key_path: str,
                                key_password: str | None = None, signature_image_path: str | None = None,
                                page_number: int = 1, box: tuple[int, int, int, int] = (200, 600, 400, 660),
                                reason: str | None = None, location: str | None = None,
                                field_name: str = "Signature1") -> None:
    """Cryptographically sign a PDF using a PAdES-style digital signature (via pyHanko).
    Unlike stamp_signature_on_pdf (which just draws an image with no security guarantee),
    this embeds a real cryptographic signature — any change to the document afterward will
    make PDF viewers (Adobe Acrobat, etc.) show the signature as invalid/broken.

    If signature_image_path is given, that image is used as the signature's visible
    appearance (in the same box a normal ink signature would occupy).
    box is (left, bottom, right, top) in points, on the given page_number (1-indexed).

    Requires the optional 'pyhanko' package: pip install pyhanko
    """
    try:
        from pyhanko import stamp
        from pyhanko.pdf_utils import images
        from pyhanko.pdf_utils.incremental_writer import IncrementalPdfFileWriter
        from pyhanko.sign import fields, signers
    except ImportError:
        raise ImportError(
            "Cryptographic PDF signing requires pyhanko. Install it with: pip install pyhanko"
        )

    signer = signers.SimpleSigner.load(
        key_file=key_path,
        cert_file=cert_path,
        key_passphrase=key_password.encode() if key_password else None,
    )

    with open(pdf_path, "rb") as inf:
        w = IncrementalPdfFileWriter(inf)
        fields.append_signature_field(
            w,
            sig_field_spec=fields.SigFieldSpec(field_name, on_page=page_number - 1, box=box),
        )

        meta = signers.PdfSignatureMetadata(field_name=field_name, reason=reason, location=location)

        stamp_style = None
        if signature_image_path:
            stamp_style = stamp.TextStampStyle(
                stamp_text="",  # empty text: the image alone becomes the visible signature
                background=images.PdfImage(signature_image_path),
            )

        pdf_signer = signers.PdfSigner(meta, signer=signer, stamp_style=stamp_style)

        with open(output_path, "wb") as outf:
            pdf_signer.sign_pdf(w, output=outf)