"""Video utilities: transcription, audio extraction, metadata, frame capture, 
and speaker diarization — powered by PyAV, faster-whisper, and Pyannote.

Requires the optional 'av' package: pip install av
(This is already installed automatically as a dependency of faster-whisper.)

Supports common video containers: MP4, MOV, MKV, AVI, WEBM, and more — since PyAV
uses FFmpeg's decoding engine internally, the same one that powers audiotools.
"""
import os
import shutil
import tempfile
import random
import subprocess
import requests


def get_video_info(path: str) -> dict:
    """Return basic metadata about a video file: duration (seconds), width, height, and frame rate."""
    import av
    container = av.open(path)
    stream = container.streams.video[0]
    duration = float(container.duration / av.time_base) if container.duration else None
    info = {
        "duration_seconds": duration,
        "width": stream.width,
        "height": stream.height,
        "fps": float(stream.average_rate) if stream.average_rate else None,
    }
    container.close()
    return info


def get_video_duration(path: str) -> float:
    """Return the duration of a video file in seconds."""
    return get_video_info(path)["duration_seconds"]


def extract_audio_from_video(video_path: str, output_audio_path: str) -> None:
    """Extract the audio track from a video file and save it as a standalone audio file
    (format is inferred from output_audio_path's extension, e.g. .mp3, .wav, .m4a)."""
    import av

    input_container = av.open(video_path)
    audio_stream = input_container.streams.audio[0]

    output_container = av.open(output_audio_path, mode="w")
    output_stream = output_container.add_stream("aac" if output_audio_path.endswith((".m4a", ".mp4")) else "mp3")

    for frame in input_container.decode(audio_stream):
        for packet in output_stream.encode(frame):
            output_container.mux(packet)

    for packet in output_stream.encode(None):
        output_container.mux(packet)

    output_container.close()
    input_container.close()


def extract_video_frame(video_path: str, timestamp_seconds: float, output_image_path: str) -> None:
    """Extract a single frame from a video at the given timestamp (seconds) and save it as an image."""
    import av

    container = av.open(video_path)
    stream = container.streams.video[0]

    target_pts = int(timestamp_seconds / stream.time_base)
    container.seek(target_pts, stream=stream)

    for frame in container.decode(stream):
        if frame.time >= timestamp_seconds:
            frame.to_image().save(output_image_path)
            break

    container.close()


def transcribe_video(path: str, model_size: str = "base", language: str | None = None) -> str:
    """Transcribe the spoken audio in a video file to plain text. Works directly on
    video containers (MP4, MOV, MKV, etc.) — the audio track is extracted automatically."""
    from quicktools.audiotools import transcribe_audio
    return transcribe_audio(path, model_size=model_size, language=language)


def transcribe_video_with_timestamps(path: str, model_size: str = "base") -> list[dict]:
    """Transcribe a video's audio into timestamped segments, each with 'start', 'end', and 'text'."""
    from quicktools.audiotools import transcribe_audio_with_timestamps
    return transcribe_audio_with_timestamps(path, model_size=model_size)


def transcribe_video_word_level(path: str, model_size: str = "base") -> list[dict]:
    """Transcribe a video's audio into word-by-word timestamps, each with 'word', 'start', and 'end'.
    Useful for generating captions synced precisely to speech."""
    from quicktools.audiotools import transcribe_audio_word_level
    return transcribe_audio_word_level(path, model_size=model_size)


# --- NEW DIARIZATION AND STYLING FEATURES ---

from quicktools import audiotools

def transcribe_video_with_speakers(path_or_url: str, hf_token: str, model_size: str = "base", device: str = "auto") -> list[dict]:
    """
    Extracts audio from a local video file or web URL (YouTube, TikTok, IG, X), 
    transcribes it, and maps the text to specific speakers.
    """
    print(f"Processing video source: {path_or_url}")
    return audiotools.transcribe_with_speakers(path_or_url, hf_token, model_size, device)


def save_video_script_to_docx(transcript_data: list[dict], output_path: str, title_text: str = "Video Script & Transcript") -> None:
    """
    Formats speaker-mapped video transcript data as a professional script in Word (.docx).
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise ImportError("Saving to Word requires python-docx. Run: pip install python-docx")

    doc = Document()
    title = doc.add_heading(title_text, level=1)
    title.alignment = 1  # Center align

    last_speaker = None

    for entry in transcript_data:
        p = doc.add_paragraph()

        # Format timestamp as [MM:SS]
        start_m, start_s = divmod(int(entry["start"]), 60)
        time_str = f"[{start_m:02d}:{start_s:02d}]"

        # Speaker header line if speaker changed
        if entry["speaker"] != last_speaker:
            speaker_run = p.add_run(f"{time_str} {entry['speaker']}:\n")
            speaker_run.bold = True
            speaker_run.font.color.rgb = RGBColor(112, 48, 160)  # Distinct Purple highlight for video speakers
            last_speaker = entry["speaker"]

        p.add_run(entry["text"])
        p.paragraph_format.space_after = Pt(8)

    doc.save(output_path)



def download_video(url: str, output_dir: str = ".", filename: str | None = None, resolution: str = "best", cookiefile: str | None = None) -> str:
    try:
        import yt_dlp
    except ImportError:
        raise ImportError("Downloading videos requires 'yt-dlp'. Install it with: pip install yt-dlp")

    os.makedirs(output_dir, exist_ok=True)
    out_template = f"{filename}.%(ext)s" if filename else "%(title)s.%(ext)s"
    target_path = os.path.join(output_dir, out_template)

    # 1. RESOLVE SHORTLINKS & NORMALIZE X.COM
    if "x.com" in url:
        url = url.replace("x.com", "twitter.com")

    final_url = url
    if "vt.tiktok.com" in url or "vm.tiktok.com" in url:
        try:
            resp = requests.head(url, allow_redirects=True, timeout=8, headers={'User-Agent': 'Mozilla/5.0'})
            final_url = resp.url
            print(f"  [quicktools] Resolved shortlink: {url} -> {final_url}")
        except Exception as e_redirect:
            print(f"  [quicktools] Shortlink resolution warning: {str(e_redirect)}")

    if not cookiefile and os.path.exists('/app/cookies.txt'):
        cookiefile = '/app/cookies.txt'

    has_ffmpeg = shutil.which("ffmpeg") is not None
    if resolution == "best":
        fmt = 'bestvideo[ext=mp4]+bestaudio[ext=m4a]/best[ext=mp4]/best' if has_ffmpeg else 'best[ext=mp4]/best'
    else:
        fmt = 'worst'

    ydl_opts = {
        'format': fmt,
        'outtmpl': target_path,
        'quiet': False,
        'no_warnings': True,
        'socket_timeout': 15,
        'merge_output_format': 'mp4' if has_ffmpeg else None,
        'http_headers': {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        },
        # === THE CORRECTED API BACKDOORS (Wrapped in Lists!) ===
        'extractor_args': {
            'youtube': {'player_client': ['android']},
            'twitter': {'api': ['syndication']},   
            'tiktok': {'app_version': ['123456']}  
        }
    }

    if has_ffmpeg:
        ydl_opts['postprocessor_args'] = ['-movflags', '+faststart']

    if cookiefile and os.path.exists(cookiefile):
        ydl_opts['cookiefile'] = cookiefile

    def attempt_download(opts: dict, target_url: str) -> str:
        with yt_dlp.YoutubeDL(opts) as ydl:
            info = ydl.extract_info(target_url, download=True)
            filepath = ydl.prepare_filename(info)
            base, _ = os.path.splitext(filepath)
            if os.path.exists(f"{base}.mp4"):
                filepath = f"{base}.mp4"
            if not os.path.exists(filepath) or os.path.getsize(filepath) == 0:
                raise RuntimeError("Download succeeded but file is missing or empty.")
            return os.path.abspath(filepath)

    # === ATTEMPT 1: Direct IPv6 ===
    print(f"  [Attempt 1] Direct IPv6 download from {final_url}...")
    ydl_opts_ipv6 = ydl_opts.copy()
    ydl_opts_ipv6['source_address'] = '0::0'
    try:
        return attempt_download(ydl_opts_ipv6, final_url)
    except Exception as e_ipv6:
        print(f"  IPv6 Attempt Failed: {str(e_ipv6)}")

    # === ATTEMPT 2: Direct IPv4 ===
    print(f"  [Attempt 2] Direct IPv4 download...")
    ydl_opts_ipv4 = ydl_opts.copy()
    try:
        return attempt_download(ydl_opts_ipv4, final_url)
    except Exception as e_ipv4:
        print(f"  IPv4 Attempt Failed: {str(e_ipv4)}")

    # === ATTEMPT 3: Rotating Proxy Endpoint ===
    proxy_url = os.getenv("RESIDENTIAL_PROXY")
    if proxy_url:
        print("  [Attempt 3] Routing through Rotating Proxy Endpoint...")
        ydl_opts_proxy = ydl_opts.copy()
        ydl_opts_proxy['proxy'] = proxy_url.strip()
        try:
            return attempt_download(ydl_opts_proxy, final_url)
        except Exception as e_proxy:
            raise RuntimeError(f"Rotating Proxy failed: {str(e_proxy)}")
    else:
        raise RuntimeError(f"Direct downloads failed (IPv6/IPv4), and no rotating proxy configured. Last error: {str(e_ipv4)}")


import subprocess
def convert_video_to_animated(input_path: str, output_path: str, target_format: str = "gif", fps: int = 15, width: int = 512) -> None:
    """Converts a video to an animated GIF or WebP sticker using FFmpeg."""
    if not shutil.which("ffmpeg"):
        raise RuntimeError("FFmpeg is required for video-to-animation conversion. Please install it.")
    
    # WhatsApp/Telegram stickers prefer a 512px boundary. We scale proportionally.
    scale_filter = f"fps={fps},scale={width}:-1:flags=lanczos"
    
    if target_format.lower() == "gif":
        # High-quality GIF generation using a 2-pass color palette
        vf_cmd = f"{scale_filter},split[s0][s1];[s0]palettegen[p];[s1][p]paletteuse"
        cmd = ["ffmpeg", "-y", "-i", input_path, "-vf", vf_cmd, "-loop", "0", output_path]
    elif target_format.lower() == "webp":
        # Force exact 512x512 transparent canvas, 15fps, and loop infinitely for WhatsApp
        wa_filter = "fps=15,scale=512:512:force_original_aspect_ratio=decrease,format=rgba,pad=512:512:(ow-iw)/2:(oh-ih)/2:color=#00000000"
        cmd = [
            "ffmpeg", "-y", "-i", input_path, 
            "-vcodec", "libwebp", 
            "-vf", wa_filter,
            "-lossless", "0", 
            "-compression_level", "6", 
            "-q:v", "40", 
            "-loop", "0", 
            "-preset", "picture", 
            "-an", "-vsync", "0", 
            output_path
        ]
    else:
        raise ValueError("Target format must be 'gif' or 'webp'")

    # Execute FFmpeg silently
    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"FFmpeg conversion failed: {result.stderr}")