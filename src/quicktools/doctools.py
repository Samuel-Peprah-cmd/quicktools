"""Word document utilities: reading, creating, editing .docx files, and converting to/from PDF."""
from docx import Document


def read_docx_text(path: str) -> str:
    """Extract and return all paragraph text from a .docx file."""
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def create_docx_from_text(text: str, output_path: str) -> None:
    """Create a new .docx file, adding each line of text as a separate paragraph."""
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(output_path)


def get_docx_word_count(path: str) -> int:
    """Count the total number of words across all paragraphs in a .docx file."""
    doc = Document(path)
    return sum(len(p.text.split()) for p in doc.paragraphs)


def get_docx_paragraph_count(path: str) -> int:
    """Count the number of paragraphs in a .docx file."""
    return len(Document(path).paragraphs)


def replace_text_in_docx(path: str, old: str, new: str, output_path: str) -> None:
    """Replace all occurrences of `old` with `new` across every paragraph, saving to output_path."""
    doc = Document(path)
    for p in doc.paragraphs:
        if old in p.text:
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
    doc.save(output_path)


def extract_docx_tables(path: str) -> list[list[list[str]]]:
    """Extract every table in a .docx file as a list of tables, each a list of rows, each a list of cell strings."""
    doc = Document(path)
    all_tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        all_tables.append(rows)
    return all_tables


def merge_docx_files(paths: list[str], output_path: str) -> None:
    """Merge multiple .docx files into one, concatenating their paragraphs in order."""
    combined = Document()
    for i, path in enumerate(paths):
        doc = Document(path)
        for p in doc.paragraphs:
            combined.add_paragraph(p.text)
        if i < len(paths) - 1:
            combined.add_page_break()
    combined.save(output_path)


def docx_to_pdf(path: str, output_path: str) -> None:
    """Convert a .docx file to PDF. Requires Microsoft Word to be installed (Windows/macOS) and
    the optional 'docx2pdf' package (pip install docx2pdf)."""
    try:
        from docx2pdf import convert
    except ImportError:
        raise ImportError(
            "docx_to_pdf() requires docx2pdf. Install it with: pip install docx2pdf\n"
            "Note: it also requires Microsoft Word to be installed on this machine."
        )
    convert(path, output_path)


def pdf_to_docx(path: str, output_path: str) -> None:
    """Convert a PDF file to .docx, preserving text layout where possible.
    Requires the optional 'pdf2docx' package (pip install pdf2docx)."""
    try:
        from pdf2docx import Converter
    except ImportError:
        raise ImportError(
            "pdf_to_docx() requires pdf2docx. Install it with: pip install pdf2docx"
        )
    cv = Converter(path)
    cv.convert(output_path)
    cv.close()

def compress_docx(input_path: str, output_path: str) -> None:
    """Compresses a Word document by applying maximum ZIP deflation to its internal structure."""
    import zipfile
    with zipfile.ZipFile(input_path, 'r') as zin:
        with zipfile.ZipFile(output_path, 'w', compression=zipfile.ZIP_DEFLATED, compresslevel=9) as zout:
            for item in zin.infolist():
                zout.writestr(item, zin.read(item.filename))